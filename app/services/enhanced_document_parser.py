"""
Enhanced document parser with layout-aware PDF -> Markdown, structured Word, Excel, Text, Markdown, and image OCR.
- PDF: MinerU (with cache) -> Docling (with cache) -> PyMuPDF fallback
- Word: Unstructured -> python-docx fallback
- Image: paddleocr (if installed) -> metadata only
"""
import logging
import re
from typing import List, Dict, Any, Optional
from pathlib import Path
import os

logger = logging.getLogger(__name__)

# MinerU support (layout-aware PDF)
try:
    from mineru import MinerU  # type: ignore
    HAS_MINERU = True
except Exception:
    MinerU = None
    HAS_MINERU = False
    logger.warning("MinerU not installed. Advanced PDF parsing will fallback to Docling/PyMuPDF.")

# Docling support (layout-aware PDF -> Markdown)
try:
    from docling.document_converter import DocumentConverter  # type: ignore
    HAS_DOCLING = True
except Exception:
    DocumentConverter = None
    HAS_DOCLING = False
    logger.warning("Docling not installed. PDF parsing will skip docling fallback.")

# PyMuPDF lightweight fallback
try:
    import fitz  # type: ignore
    HAS_PYMUPDF = True
except Exception:
    HAS_PYMUPDF = False

# Excel support
try:
    import openpyxl  # type: ignore
    HAS_EXCEL = True
except Exception:
    HAS_EXCEL = False

# Word fallbacks
try:
    from docx import Document  # type: ignore
    HAS_DOCX = True
except Exception:
    HAS_DOCX = False

try:
    from unstructured.partition.docx import partition_docx  # type: ignore
    HAS_UNSTRUCTURED = True
except Exception:
    HAS_UNSTRUCTURED = False

# Image OCR
try:
    from paddleocr import PaddleOCR  # type: ignore
    HAS_PADDLE = True
except Exception:
    PaddleOCR = None
    HAS_PADDLE = False

try:
    from PIL import Image  # type: ignore
    HAS_PIL = True
except Exception:
    HAS_PIL = False


def _ensure_dir(path: Optional[str]) -> Optional[Path]:
    if not path:
        return None
    p = Path(path)
    p.mkdir(parents=True, exist_ok=True)
    return p


class EnhancedDocumentParser:
    def __init__(self, llm_client=None, use_mineru: bool = True, use_docling: bool = True):
        self.llm_client = llm_client
        self.use_mineru = use_mineru and HAS_MINERU
        self.use_docling = use_docling and HAS_DOCLING
        self.mineru_cache = _ensure_dir(os.getenv("MINERU_CACHE_DIR"))
        self.docling_cache = _ensure_dir(os.getenv("DOCLING_CACHE_DIR"))

        self.mineru = None
        if self.use_mineru:
            try:
                self.mineru = MinerU(backend="pipeline", cache_dir=str(self.mineru_cache) if self.mineru_cache else None)  # type: ignore
                logger.info("MinerU initialized with pipeline backend")
            except Exception as e:
                logger.warning(f"Failed to initialize MinerU: {e}, falling back to Docling/PyMuPDF")
                self.use_mineru = False

        if self.use_docling and not HAS_DOCLING:
            self.use_docling = False

        self.ocr = None
        if HAS_PADDLE:
            try:
                self.ocr = PaddleOCR(use_angle_cls=True, lang="ch")  # type: ignore
            except Exception as e:
                logger.warning(f"PaddleOCR init failed: {e}")
                self.ocr = None

    def parse(self, file_path: Path) -> Dict[str, Any]:
        suffix = file_path.suffix.lower()
        if suffix == ".pdf":
            return self.parse_pdf(file_path)
        if suffix in [".xlsx", ".xls"]:
            return self.parse_excel(file_path)
        if suffix in [".docx", ".doc"]:
            return self.parse_docx(file_path)
        if suffix in [".png", ".jpg", ".jpeg", ".bmp"]:
            return self.parse_image(file_path)
        if suffix == ".txt":
            return self.parse_text(file_path)
        if suffix in [".md", ".markdown"]:
            return self.parse_markdown(file_path)
        raise ValueError(f"Unsupported file format: {suffix}")

    def parse_pdf(self, pdf_path: Path) -> Dict[str, Any]:
        if self.use_mineru and self.mineru:
            try:
                return self._parse_pdf_with_mineru(pdf_path)
            except Exception as e:
                logger.warning(f"MinerU parse failed, falling back to Docling/PyMuPDF: {e}")
        if self.use_docling and HAS_DOCLING:
            try:
                return self._parse_pdf_with_docling(pdf_path)
            except Exception as e:
                logger.warning(f"Docling parse failed, falling back to PyMuPDF: {e}")
        if HAS_PYMUPDF:
            return self._parse_pdf_with_pymupdf(pdf_path)
        raise ValueError("No PDF parser available. Install mineru or docling or pymupdf.")

    def _parse_pdf_with_mineru(self, pdf_path: Path) -> Dict[str, Any]:
        result = self.mineru.parse(str(pdf_path))  # type: ignore
        markdown = getattr(result, "markdown", "") or ""
        text = self._markdown_to_text(markdown) if markdown else ""
        if not text and hasattr(result, "text"):
            text = result.text
        tables = getattr(result, "tables", []) or []
        images = getattr(result, "images", []) or []
        return {
            "text": text,
            "markdown": markdown or text,
            "metadata": {
                "parser": "mineru",
                "format": "pdf",
                "has_tables": bool(tables),
                "has_images": bool(images),
            },
            "tables": tables,
            "images": images,
        }

    def _parse_pdf_with_docling(self, pdf_path: Path) -> Dict[str, Any]:
        converter = DocumentConverter(cache_dir=str(self.docling_cache) if self.docling_cache else None)  # type: ignore
        doc = converter.convert(pdf_path)
        markdown = doc.document.export_to_markdown()
        text = self._markdown_to_text(markdown)
        headings = len(re.findall(r"^#", markdown, flags=re.MULTILINE))
        return {
            "text": text,
            "markdown": markdown,
            "metadata": {"parser": "docling", "format": "pdf", "heading_count": headings},
            "tables": [],
            "images": [],
        }

    def _parse_pdf_with_pymupdf(self, pdf_path: Path) -> Dict[str, Any]:
        doc = fitz.open(str(pdf_path))  # type: ignore
        pages_text: List[str] = []
        for page in doc:
            pages_text.append(page.get_text())
        doc.close()
        text = "\n\n".join(pages_text)
        markdown = self._text_to_markdown_with_headers(text)
        return {
            "text": text,
            "markdown": markdown,
            "metadata": {"parser": "pymupdf", "format": "pdf", "page_count": len(pages_text)},
            "tables": [],
            "images": [],
        }

    def parse_docx(self, docx_path: Path) -> Dict[str, Any]:
        if HAS_UNSTRUCTURED:
            try:
                elements = partition_docx(filename=str(docx_path))  # type: ignore
                md_lines: List[str] = []
                for el in elements:
                    category = getattr(el, "category", "") or el.__class__.__name__
                    text = getattr(el, "text", "")
                    if not text:
                        continue
                    if category.lower() == "title":
                        md_lines.append(f"# {text}")
                    elif "heading" in category.lower():
                        md_lines.append(f"## {text}")
                    elif "list" in category.lower():
                        md_lines.append(f"- {text}")
                    else:
                        md_lines.append(text)
                markdown = "\n".join(md_lines)
                return {
                    "text": self._markdown_to_text(markdown),
                    "markdown": markdown,
                    "metadata": {"parser": "unstructured", "format": "docx", "element_count": len(md_lines)},
                    "tables": [],
                    "images": [],
                }
            except Exception as e:
                logger.warning(f"Unstructured docx parsing failed, fallback to python-docx: {e}")

        if not HAS_DOCX:
            raise ValueError("DOCX parsing requires python-docx or unstructured.")

        doc = Document(str(docx_path))  # type: ignore
        text_parts: List[str] = []
        tables: List[Dict[str, Any]] = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                if any(cell.strip() for cell in row_data):
                    table_data.append(row_data)
                    text_parts.append("\t".join(row_data))
            if table_data:
                tables.append({"data": table_data})
        content = "\n".join(text_parts)
        markdown = self._text_to_markdown_with_headers(content)
        return {
            "text": content,
            "markdown": markdown,
            "metadata": {"parser": "python-docx", "format": "docx", "paragraph_count": len([p for p in doc.paragraphs if p.text.strip()]), "table_count": len(tables)},
            "tables": tables,
            "images": [],
        }

    def parse_excel(self, excel_path: Path) -> Dict[str, Any]:
        if not HAS_EXCEL:
            raise ValueError("Excel parsing requires openpyxl.")
        wb = openpyxl.load_workbook(str(excel_path), data_only=True)  # type: ignore
        tables: List[Dict[str, Any]] = []
        for sheet in wb.sheetnames:
            ws = wb[sheet]
            sheet_data: List[List[Any]] = []  # type: ignore
            for row in ws.iter_rows(values_only=True):
                row_values = [str(cell) if cell is not None else "" for cell in row]
                if any(value.strip() for value in row_values):
                    sheet_data.append(row_values)
            if sheet_data:
                tables.append({"sheet": sheet, "data": sheet_data})
        combined_text = "\n".join(["\t".join(row) for table in tables for row in table.get("data", [])])
        return {
            "text": combined_text,
            "markdown": combined_text,
            "metadata": {"parser": "openpyxl", "format": "excel", "sheet_count": len(wb.sheetnames), "table_count": len(tables)},
            "tables": tables,
            "images": [],
        }

    def parse_image(self, img_path: Path) -> Dict[str, Any]:
        text = ""
        if self.ocr and HAS_PIL:
            try:
                img = Image.open(img_path)  # type: ignore
                result = self.ocr.ocr(img, cls=True)
                lines = []
                for line in result:
                    for item in line:
                        if len(item) >= 2 and item[1]:
                            lines.append(item[1][0])
                text = "\n".join(lines)
            except Exception as e:
                logger.warning(f"OCR failed for {img_path.name}: {e}")
        return {
            "text": text,
            "markdown": text,
            "metadata": {"parser": "paddleocr" if text else "image-metadata", "format": "image", "filename": img_path.name},
            "tables": [],
            "images": [str(img_path)],
        }

    def parse_text(self, file_path: Path) -> Dict[str, Any]:
        content = file_path.read_text(encoding="utf-8")
        return {
            "text": content,
            "markdown": content,
            "metadata": {"parser": "text", "format": "txt", "line_count": len(content.splitlines())},
            "tables": [],
            "images": [],
        }

    def parse_markdown(self, file_path: Path) -> Dict[str, Any]:
        content = file_path.read_text(encoding="utf-8")
        return {
            "text": self._markdown_to_text(content),
            "markdown": content,
            "metadata": {"parser": "markdown", "format": "md"},
            "tables": [],
            "images": [],
        }

    def _markdown_to_text(self, markdown: str) -> str:
        text = re.sub(r"#{1,6}\s+", "", markdown)
        text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
        text = re.sub(r"\*(.+?)\*", r"\1", text)
        text = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", text)
        text = re.sub(r"```[\s\S]*?```", "", text)
        text = re.sub(r"`(.+?)`", r"\1", text)
        return text

    def _text_to_markdown_with_headers(self, text: str) -> str:
        md_lines: List[str] = []
        for line in text.splitlines():
            stripped = line.strip()
            if not stripped:
                md_lines.append("")
                continue
            if len(stripped) < 80 and stripped.isupper():
                md_lines.append(f"## {stripped.title()}")
            elif stripped.endswith(":"):
                md_lines.append(f"## {stripped}")
            else:
                md_lines.append(stripped)
        return "\n".join(md_lines)
