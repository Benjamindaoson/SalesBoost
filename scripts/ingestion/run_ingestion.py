"""
SalesBoost RAG Data Ingestion Pipeline
数据处理与注入流水线 (RAG + GraphRAG)

功能：
1. 遍历加载本地文档 (.docx, .pdf, .txt)
2. 切片 (Chunking) 与 清洗
3. 向量化并存入 Vector DB (KnowledgeService)
4. LLM 自动化提取实体关系并存入 GraphRAG
5. 持久化数据

Usage:
python scripts/run_ingestion.py --dir "path/to/documents"
"""
import argparse
import asyncio
import logging
import json
from pathlib import Path
from typing import List, Dict, Any
from tqdm import tqdm

# Custom Loaders to avoid langchain_community dependency issues
import docx
from pypdf import PdfReader

import hashlib
from cognitive.tools.connectors.ingestion.semantic_chunker import SemanticChunker

# Simple Document Class
class Document:
    def __init__(self, page_content: str, metadata: Dict[str, Any] = None):
        self.page_content = page_content
        self.metadata = metadata or {}

# 服务
from cognitive.skills.study.knowledge_service import KnowledgeService
from cognitive.skills.study.advanced_rag.graph_rag import GraphRAGService
from cognitive.infra.gateway.model_gateway import ModelGateway
from core.config import get_settings

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

class DataIngestionPipeline:
    def __init__(self, data_dir: str, enable_graph_extraction: bool = True):
        self.data_dir = Path(data_dir)
        self.enable_graph_extraction = enable_graph_extraction
        self.settings = get_settings()
        
        # 初始化服务
        self.knowledge_service = KnowledgeService()
        self.graph_rag_service = GraphRAGService()
        self.model_gateway = ModelGateway()
        
        # 文本切分器
        self.semantic_chunker = SemanticChunker(max_tokens=500)

    def _load_txt(self, file_path: Path) -> List[Document]:
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
            return [Document(page_content=text, metadata={"source": file_path.name})]
        except Exception as e:
            logger.warning(f"Error loading TXT {file_path}: {e}")
            return []

    def _load_pdf(self, file_path: Path) -> List[Document]:
        docs = []
        try:
            reader = PdfReader(str(file_path))
            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                if text:
                    docs.append(Document(page_content=text, metadata={"source": file_path.name, "page": i}))
        except Exception as e:
            logger.warning(f"Error loading PDF {file_path}: {e}")
        return docs

    def _load_docx(self, file_path: Path) -> List[Document]:
        try:
            doc = docx.Document(str(file_path))
            text = "\n".join([para.text for para in doc.paragraphs])
            return [Document(page_content=text, metadata={"source": file_path.name})]
        except Exception as e:
            logger.warning(f"Error loading DOCX {file_path}: {e}")
            return []

    def load_documents(self) -> List[Document]:
        """Phase 1: 加载与清洗"""
        logger.info(f"Scanning directory: {self.data_dir}")
        documents = []
        
        if not self.data_dir.exists():
            logger.error(f"Directory not found: {self.data_dir}")
            return []

        # 遍历文件
        files = []
        for ext in ["*.txt", "*.pdf", "*.docx"]:
            # Filter out hidden/metadata files
            candidates = list(self.data_dir.rglob(ext))
            files.extend([f for f in candidates if not f.name.startswith("._") and not f.name.startswith("~")])
            
        logger.info(f"Found {len(files)} files")
        
        for file_path in tqdm(files, desc="Loading files"):
            try:
                docs = []
                if file_path.suffix.lower() == ".txt":
                    docs = self._load_txt(file_path)
                elif file_path.suffix.lower() == ".pdf":
                    docs = self._load_pdf(file_path)
                elif file_path.suffix.lower() == ".docx":
                    docs = self._load_docx(file_path)
                
                for doc in docs:
                    # 清洗
                    content = doc.page_content
                    content = content.replace("\x00", "") # 去除空字符
                    # 去除多余换行
                    content = "\n".join([line.strip() for line in content.split("\n") if line.strip()])
                    
                    if len(content) < 10:
                        continue
                        
                    doc.page_content = content
                    # Ensure metadata has source
                    if "source" not in doc.metadata:
                        doc.metadata["source"] = file_path.name
                        
                    documents.append(doc)
                        
            except Exception as e:
                logger.warning(f"Failed to load {file_path}: {e}")
                continue
                
        logger.info(f"Loaded {len(documents)} valid documents")
        return documents

    def chunk_documents(self, documents: List[Document]) -> List[Document]:
        """Phase 2: Semantic Chunking (Deterministic)"""
        logger.info("Chunking documents (Semantic + Structural + Deterministic)...")
        final_docs = []
        
        for doc in tqdm(documents, desc="Semantic Chunking"):
            # Calculate Doc SHA256 for stable ID
            doc_content_bytes = doc.page_content.encode('utf-8')
            doc_sha256 = hashlib.sha256(doc_content_bytes).hexdigest()
            
            # Use source from metadata or fallback
            source_id = doc.metadata.get("source", "unknown")
            
            # No async needed now as it is deterministic
            chunks_data = self.semantic_chunker.chunk_document(
                text=doc.page_content, 
                source_id=source_id,
                doc_sha256=doc_sha256,
                base_metadata=doc.metadata
            )
            
            for chunk_data in chunks_data:
                content = chunk_data.pop("content")
                new_doc = Document(page_content=content, metadata=chunk_data)
                final_docs.append(new_doc)
        
        logger.info(f"Created {len(final_docs)} semantic chunks")
        return final_docs

    async def ingest_vectors(self, chunks: List[Document]):
        """Phase 2: 向量化并存入 ChromaDB"""
        logger.info("Ingesting into Vector DB...")
        
        # 批量处理以避免内存溢出
        batch_size = 100
        for i in tqdm(range(0, len(chunks), batch_size), desc="Vectorizing"):
            batch = chunks[i:i+batch_size]
            
            texts = [doc.page_content for doc in batch]
            metadatas = [doc.metadata for doc in batch]
            
            # 模拟批量添加 (KnowledgeService 需要实现 add_documents)
            try:
                if self.knowledge_service.collection:
                    self.knowledge_service.collection.add(
                        documents=texts,
                        metadatas=metadatas,
                        ids=[f"chunk_{i}_{j}" for j in range(len(batch))]
                    )
                else:
                    logger.warning("KnowledgeService collection not initialized, skipping batch ingestion")
            except Exception as e:
                logger.error(f"Vector ingestion failed for batch {i}: {e}")

    async def extract_graph_data(self, chunk: Document):
        """使用 LLM 提取图谱数据"""
        prompt = f"""
        你是一个知识图谱专家。请从以下文本中提取实体和关系。
        
        【目标实体类型】
        - Product (产品名)
        - Benefit (权益/优势)
        - Objection (客户异议/问题)
        
        【目标关系类型】
        - HAS_BENEFIT (Product -> Benefit)
        - SOLVED_BY (Objection -> Benefit)
        - RELATED_TO (Objection -> Product)
        
        【文本】
        {chunk.page_content}
        
        【输出格式】
        请仅输出合法的 JSON，不要包含 Markdown 标记：
        {{
            "nodes": [
                {{"id": "实体名", "type": "类型", "content": "描述"}}, ...
            ],
            "edges": [
                {{"source": "实体名", "target": "实体名", "relation": "关系类型"}}, ...
            ]
        }}
        """
        
        try:
            response = await self.model_gateway.chat_completion(
                messages=[{"role": "user", "content": prompt}],
                model="gpt-3.5-turbo", # 使用快速模型
                temperature=0.0
            )
            content = response.choices[0].message.content
            # 清理 Markdown
            if "```json" in content:
                content = content.split("```json")[1].split("```")[0]
            elif "```" in content:
                content = content.split("```")[1].split("```")[0]
                
            data = json.loads(content.strip())
            return data
        except Exception:
            # logger.warning(f"Graph extraction failed: {e}")
            return None

    async def ingest_graph(self, chunks: List[Document]):
        """Phase 3: 自动化图谱构建"""
        if not self.enable_graph_extraction:
            logger.info("Graph extraction disabled, skipping.")
            return

        logger.info("Extracting Graph Data (This may take a while)...")
        
        # 并发控制
        sem = asyncio.Semaphore(5) # 限制并发数为 5
        
        async def process_chunk(chunk):
            async with sem:
                data = await self.extract_graph_data(chunk)
                if data:
                    # 写入图谱服务
                    for node in data.get("nodes", []):
                        self.graph_rag_service.add_node(
                            id=node["id"],
                            type=node["type"],
                            content=node.get("content", node["id"]),
                            metadata={"source": chunk.metadata.get("source")}
                        )
                    for edge in data.get("edges", []):
                        self.graph_rag_service.add_edge(
                            source=edge["source"],
                            target=edge["target"],
                            relation=edge["relation"]
                        )
                        
        tasks = [process_chunk(chunk) for chunk in chunks[:50]] # 限制前50个chunk做演示，防止token爆炸
        
        # 显示进度
        for f in tqdm(asyncio.as_completed(tasks), total=len(tasks), desc="Building Graph"):
            await f

    def save_data(self):
        """Phase 4: 持久化"""
        logger.info("Persisting data...")
        
        # 1. Vector DB (Chroma 自动持久化)
        # 2. Graph Data (Pickle dump)
        import pickle
        graph_dir = Path("./data/graph_store")
        graph_dir.mkdir(parents=True, exist_ok=True)
        
        try:
            with open(graph_dir / "nodes.pkl", "wb") as f:
                pickle.dump(self.graph_rag_service.nodes, f)
            with open(graph_dir / "edges.pkl", "wb") as f:
                pickle.dump(self.graph_rag_service.edges, f)
            with open(graph_dir / "adjacency.pkl", "wb") as f:
                pickle.dump(self.graph_rag_service.adjacency, f)
            logger.info(f"Graph data saved to {graph_dir}")
        except Exception as e:
            logger.error(f"Failed to save graph data: {e}")

async def main():
    parser = argparse.ArgumentParser(description="SalesBoost Data Ingestion")
    parser.add_argument("--dir", type=str, default="data", help="Data directory path")
    parser.add_argument("--no-graph", action="store_true", help="Disable graph extraction")
    
    args = parser.parse_args()
    
    pipeline = DataIngestionPipeline(
        data_dir=args.dir,
        enable_graph_extraction=not args.no_graph
    )
    
    # 1. Load
    docs = pipeline.load_documents()
    if not docs:
        logger.warning("No documents loaded. Exiting.")
        return
        
    # 2. Chunk
    chunks = pipeline.chunk_documents(docs)
    
    # 3. Vector Ingestion
    await pipeline.ingest_vectors(chunks)
    
    # 4. Graph Ingestion
    await pipeline.ingest_graph(chunks)
    
    # 5. Save
    pipeline.save_data()
    
    logger.info("Ingestion Pipeline Completed Successfully!")

if __name__ == "__main__":
    asyncio.run(main())
