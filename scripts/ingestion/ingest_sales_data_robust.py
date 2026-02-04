"""
销售数据摄入脚本（健壮版本）

处理各种错误情况，即使某些模块失败也能继续处理其他文件。
"""
import asyncio
import logging
import sys
from pathlib import Path
from typing import Dict, Any, Optional
import uuid

# 添加项目根目录到Python路径
project_root = Path(__file__).parent.parent
sys.path.insert(0, str(project_root))

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# 数据目录
DATA_DIR = Path(__file__).parent.parent / "销冠能力复制数据库"

# 检查依赖
HAS_EXCEL = False
HAS_DOCX = False
HAS_PDF = False
HAS_MINERU = False

try:
    import openpyxl
    HAS_EXCEL = True
except ImportError:
    logger.warning("openpyxl not installed")

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    logger.warning("python-docx not installed")

try:
    import fitz
    HAS_PDF = True
except ImportError:
    logger.warning("PyMuPDF not installed")

try:
    from mineru import MinerU
    HAS_MINERU = True
except ImportError:
    logger.warning("MinerU not installed")


class RobustSalesDataIngester:
    """健壮的销售数据摄入器"""
    
    def __init__(self, org_id: Optional[str] = None):
        """初始化"""
        self.org_id = org_id or "public"
        self.stats = {
            "excel_files": 0,
            "pdf_files": 0,
            "docx_files": 0,
            "ppt_files": 0,
            "total_documents": 0,
            "total_entities": 0,
            "total_relations": 0,
            "errors": [],
            "skipped": [],
        }
        
        # 延迟初始化服务（避免启动时失败）
        self.knowledge_service = None
        self.graph_rag_service = None
        self.enhanced_parser = None
    
    def _init_services(self):
        """延迟初始化服务"""
        if self.knowledge_service is None:
            try:
                from cognitive.skills.study.knowledge_service import KnowledgeService
                self.knowledge_service = KnowledgeService(org_id=self.org_id)
                logger.info("KnowledgeService initialized")
            except Exception as e:
                logger.error(f"Failed to initialize KnowledgeService: {e}")
                self.stats["errors"].append(f"KnowledgeService init: {str(e)}")
        
        if self.graph_rag_service is None:
            try:
                from cognitive.skills.study.graph_rag_service import GraphRAGService
                self.graph_rag_service = GraphRAGService(org_id=self.org_id)
                logger.info("GraphRAGService initialized")
            except Exception as e:
                logger.warning(f"Failed to initialize GraphRAGService: {e}")
                self.stats["errors"].append(f"GraphRAGService init: {str(e)}")
        
        if self.enhanced_parser is None:
            try:
                from cognitive.tools.parsers.enhanced_parser import EnhancedDocumentParser
                self.enhanced_parser = EnhancedDocumentParser(use_mineru=HAS_MINERU)
                logger.info("EnhancedDocumentParser initialized")
            except Exception as e:
                logger.warning(f"Failed to initialize EnhancedDocumentParser: {e}")
    
    async def ingest_all(self) -> Dict[str, Any]:
        """摄入所有数据"""
        logger.info(f"Starting data ingestion from: {DATA_DIR}")
        
        if not DATA_DIR.exists():
            logger.error(f"Data directory not found: {DATA_DIR}")
            return self.stats
        
        # 初始化服务
        self._init_services()
        
        # 如果服务都失败，至少可以解析文件
        if not self.knowledge_service and not self.graph_rag_service:
            logger.error("Both KnowledgeService and GraphRAGService failed to initialize")
            logger.info("Will parse files but not ingest to database")
        
        # 1. 摄入产品权益
        await self._ingest_product_benefits()
        
        # 2. 摄入SOP和话术
        await self._ingest_sop_and_scripts()
        
        # 3. 摄入销售冠军经验
        await self._ingest_champion_experience()
        
        logger.info(f"Ingestion completed. Stats: {self.stats}")
        return self.stats
    
    async def _ingest_product_benefits(self):
        """摄入产品权益"""
        product_dir = DATA_DIR / "产品权益"
        if not product_dir.exists():
            logger.warning("Product benefits directory not found")
            return
        
        logger.info("Ingesting product benefits...")
        
        for excel_file in product_dir.glob("*.xlsx"):
            if excel_file.name.startswith("~$"):
                continue
            try:
                await self._ingest_excel_file(excel_file, doc_type="product_benefit")
                self.stats["excel_files"] += 1
            except Exception as e:
                logger.error(f"Failed to ingest {excel_file.name}: {e}")
                self.stats["errors"].append(f"{excel_file.name}: {str(e)}")
        
        # 竞品分析
        competitor_dir = product_dir / "竞品分析"
        if competitor_dir.exists():
            for excel_file in competitor_dir.glob("*.xlsx"):
                if excel_file.name.startswith("~$"):
                    continue
                try:
                    await self._ingest_excel_file(excel_file, doc_type="competitor_analysis")
                    self.stats["excel_files"] += 1
                except Exception as e:
                    logger.error(f"Failed to ingest {excel_file.name}: {e}")
                    self.stats["errors"].append(f"{excel_file.name}: {str(e)}")
    
    async def _ingest_sop_and_scripts(self):
        """摄入SOP和话术"""
        sop_dir = DATA_DIR / "销售成交营销SOP和话术"
        if not sop_dir.exists():
            logger.warning("SOP directory not found")
            return
        
        logger.info("Ingesting SOP and scripts...")
        
        # PDF
        for pdf_file in sop_dir.glob("*.pdf"):
            try:
                await self._ingest_pdf_file(pdf_file, doc_type="sop")
                self.stats["pdf_files"] += 1
            except Exception as e:
                logger.error(f"Failed to ingest {pdf_file.name}: {e}")
                self.stats["errors"].append(f"{pdf_file.name}: {str(e)}")
        
        # Word
        for docx_file in sop_dir.glob("*.docx"):
            try:
                await self._ingest_docx_file(docx_file, doc_type="script")
                self.stats["docx_files"] += 1
            except Exception as e:
                logger.error(f"Failed to ingest {docx_file.name}: {e}")
                self.stats["errors"].append(f"{docx_file.name}: {str(e)}")
    
    async def _ingest_champion_experience(self):
        """摄入销售冠军经验"""
        experience_dir = DATA_DIR / "销售冠军成交经验分享"
        if not experience_dir.exists():
            logger.warning("Experience directory not found")
            return
        
        logger.info("Ingesting champion experience...")
        
        for docx_file in experience_dir.glob("*.docx"):
            try:
                await self._ingest_docx_file(docx_file, doc_type="case")
                self.stats["docx_files"] += 1
            except Exception as e:
                logger.error(f"Failed to ingest {docx_file.name}: {e}")
                self.stats["errors"].append(f"{docx_file.name}: {str(e)}")
    
    async def _ingest_excel_file(self, file_path: Path, doc_type: str):
        """摄入Excel文件"""
        if not HAS_EXCEL:
            self.stats["skipped"].append(f"{file_path.name}: Excel support not available")
            return
        
        logger.info(f"Processing Excel: {file_path.name}")
        
        content = None
        metadata = {}
        
        # 优先使用增强解析器
        if self.enhanced_parser:
            try:
                parsed = self.enhanced_parser.parse_excel(file_path)
                content = parsed["text"]
                metadata = parsed["metadata"]
                logger.info(f"Excel parsed with enhanced parser: {file_path.name}")
            except Exception as e:
                logger.warning(f"Enhanced parser failed: {e}, trying openpyxl directly")
        
        # 降级到直接使用openpyxl
        if not content:
            try:
                import openpyxl
                workbook = openpyxl.load_workbook(file_path, data_only=True)
                text_parts = []
                
                for sheet_name in workbook.sheetnames:
                    sheet = workbook[sheet_name]
                    text_parts.append(f"【{sheet_name}】\n")
                    
                    for row in sheet.iter_rows(values_only=True):
                        row_text = "\t".join(str(cell) if cell else "" for cell in row)
                        if row_text.strip():
                            text_parts.append(row_text)
                    text_parts.append("\n")
                
                content = "\n".join(text_parts)
                metadata = {"file_type": "excel", "sheet_count": len(workbook.sheetnames), "parser": "openpyxl"}
                logger.info(f"Excel parsed with openpyxl: {file_path.name}")
            except Exception as e:
                # 如果openpyxl也失败，尝试pandas
                try:
                    import pandas as pd
                    xls = pd.ExcelFile(file_path, engine='openpyxl')
                    text_parts = []
                    
                    for sheet_name in xls.sheet_names:
                        df = pd.read_excel(xls, sheet_name=sheet_name, engine='openpyxl')
                        text_parts.append(f"【{sheet_name}】\n")
                        text_parts.append(df.to_string(index=False))
                        text_parts.append("\n")
                    
                    content = "\n".join(text_parts)
                    metadata = {"file_type": "excel", "sheet_count": len(xls.sheet_names), "parser": "pandas"}
                    logger.info(f"Excel parsed with pandas: {file_path.name}")
                except Exception as e2:
                    logger.error(f"All Excel parsers failed for {file_path.name}: {e}, {e2}")
                    self.stats["errors"].append(f"{file_path.name}: Excel parsing failed - {str(e2)}")
                    return
        
        if not content or not content.strip():
            logger.warning(f"Empty Excel file: {file_path.name}")
            return
        
        await self._ingest_content(
            content=content,
            filename=file_path.name,
            doc_type=doc_type,
            metadata=metadata,
        )
    
    async def _ingest_pdf_file(self, file_path: Path, doc_type: str):
        """摄入PDF文件"""
        logger.info(f"Processing PDF: {file_path.name}")
        
        content = None
        parser_used = None
        
        # 优先使用MinerU（如果可用）
        if self.enhanced_parser and HAS_MINERU:
            try:
                parsed = self.enhanced_parser.parse_pdf(file_path)
                content = parsed["text"]
                parser_used = "mineru"
                logger.info(f"PDF parsed with MinerU: {file_path.name}")
            except Exception as e:
                logger.warning(f"MinerU parsing failed: {e}")
        
        # 如果MinerU失败，尝试PyMuPDF
        if not content:
            # 检查PyMuPDF是否可用（避免DLL错误）
            try:
                import fitz
                pdf_doc = fitz.open(str(file_path))
                text_parts = []
                for page_num in range(len(pdf_doc)):
                    page = pdf_doc[page_num]
                    text = page.get_text()
                    if text.strip():
                        text_parts.append(text)
                pdf_doc.close()
                content = "\n\n".join(text_parts)
                parser_used = "pymupdf"
                logger.info(f"PDF parsed with PyMuPDF: {file_path.name}")
            except Exception as e:
                logger.warning(f"PyMuPDF parsing failed: {e}")
                # 如果都失败，记录错误但继续
                self.stats["errors"].append(f"{file_path.name}: PDF parsing failed - {str(e)}")
                self.stats["skipped"].append(f"{file_path.name}: No PDF parser available")
                return
        
        if not content or not content.strip():
            logger.warning(f"Empty PDF file: {file_path.name}")
            return
        
        await self._ingest_content(
            content=content,
            filename=file_path.name,
            doc_type=doc_type,
            metadata={"file_type": "pdf", "parser": parser_used},
        )
    
    async def _ingest_docx_file(self, file_path: Path, doc_type: str):
        """摄入Word文档"""
        if not HAS_DOCX:
            self.stats["skipped"].append(f"{file_path.name}: DOCX support not available")
            return
        
        logger.info(f"Processing DOCX: {file_path.name}")
        
        try:
            from docx import Document
            doc = Document(str(file_path))
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            for table in doc.tables:
                for row in table.rows:
                    row_text = "\t".join(cell.text for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)
            
            content = "\n".join(text_parts)
            
            if not content.strip():
                logger.warning(f"Empty DOCX file: {file_path.name}")
                return
            
            await self._ingest_content(
                content=content,
                filename=file_path.name,
                doc_type=doc_type,
                metadata={"file_type": "docx"},
            )
            
        except Exception as e:
            raise Exception(f"DOCX parsing error: {e}")
    
    async def _ingest_content(
        self,
        content: str,
        filename: str,
        doc_type: str,
        metadata: Dict[str, Any],
    ):
        """将内容摄入到系统"""
        doc_id = f"{doc_type}_{uuid.uuid4().hex[:8]}"
        
        # 摄入到向量数据库
        if self.knowledge_service:
            try:
                # 对于Excel文件，使用text/plain，但标记为excel类型
                content_type_map = {
                    "excel": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    "pdf": "application/pdf",
                }
                content_type = content_type_map.get(metadata.get("file_type", ""), "text/plain")
                
                await self.knowledge_service.add_document_with_processing(
                    content=content.encode("utf-8"),
                    filename=filename,
                    content_type=content_type,
                    meta={
                        **metadata,
                        "doc_id": doc_id,
                        "type": doc_type,
                        "org_id": self.org_id,
                    },
                    doc_type=doc_type,
                )
                logger.info(f"Added to vector DB: {filename}")
            except Exception as e:
                logger.error(f"Failed to add to vector DB: {e}")
                self.stats["errors"].append(f"Vector DB {filename}: {str(e)}")
                # 不抛出异常，继续处理GraphRAG
        
        # 摄入到GraphRAG
        if self.graph_rag_service:
            try:
                result = await self.graph_rag_service.ingest_document(
                    doc_id=doc_id,
                    text=content,
                    metadata={
                        **metadata,
                        "doc_type": doc_type,
                        "org_id": self.org_id,
                    },
                )
                self.stats["total_entities"] += result.get("total_entities", 0)
                self.stats["total_relations"] += result.get("total_triples", 0)
                logger.info(f"Added to GraphRAG: {filename}")
            except Exception as e:
                logger.error(f"Failed to add to GraphRAG: {e}")
                self.stats["errors"].append(f"GraphRAG {filename}: {str(e)}")
        
        self.stats["total_documents"] += 1
    
    def print_statistics(self):
        """打印统计信息"""
        print("\n" + "="*60)
        print("数据摄入统计")
        print("="*60)
        print(f"Excel文件: {self.stats['excel_files']}")
        print(f"PDF文件: {self.stats['pdf_files']}")
        print(f"Word文件: {self.stats['docx_files']}")
        print(f"总文档数: {self.stats['total_documents']}")
        print(f"总实体数: {self.stats['total_entities']}")
        print(f"总关系数: {self.stats['total_relations']}")
        print(f"跳过文件: {len(self.stats['skipped'])}")
        print(f"错误数: {len(self.stats['errors'])}")
        
        if self.stats['errors']:
            print("\n错误列表（前10个）:")
            for error in self.stats['errors'][:10]:
                print(f"  - {error}")
        
        print("="*60 + "\n")


async def main():
    """主函数"""
    ingester = RobustSalesDataIngester(org_id="public")
    
    try:
        await ingester.ingest_all()
        ingester.print_statistics()
    except Exception as e:
        logger.error(f"Data ingestion failed: {e}", exc_info=True)
        ingester.print_statistics()
        raise


if __name__ == "__main__":
    asyncio.run(main())

