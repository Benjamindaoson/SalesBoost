"""
销售数据摄入脚本

将"销冠能力复制数据库"文件夹中的数据导入到系统中：
1. RAG系统（向量数据库）
2. GraphRAG系统（知识图谱）
3. 多智能体系统（话术、案例、策略）

数据来源：
- 产品权益（Excel）
- 销售成交营销SOP和话术（PDF、Word、PPT）
- 销售冠军成交经验分享（Word）
- 销售录音（音频，可选，需要STT）
"""
import asyncio
import logging
import os
import sys
from pathlib import Path
from typing import List, Dict, Any, Optional
import uuid

# 添加项目根目录到Python路径
project_root = Path(__file__).parent.parent
sys.path.insert(0, str(project_root))

# 文档处理
try:
    import fitz  # PyMuPDF
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

try:
    import openpyxl
    HAS_EXCEL = True
except ImportError:
    HAS_EXCEL = False

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# 系统导入
from app.services.knowledge_service import KnowledgeService
from app.services.graph_rag_service import GraphRAGService
from app.services.document_parser import DocumentProcessor
from app.services.enhanced_document_parser import EnhancedDocumentParser
from app.schemas.fsm import SalesStage

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# 数据目录
DATA_DIR = Path(__file__).parent.parent / "销冠能力复制数据库"


class SalesDataIngester:
    """销售数据摄入器"""
    
    def __init__(self, org_id: Optional[str] = None):
        """
        初始化数据摄入器
        
        Args:
            org_id: 组织ID（默认使用"public"）
        """
        self.org_id = org_id or "public"
        
        # 初始化服务
        self.knowledge_service = KnowledgeService(org_id=self.org_id)
        self.graph_rag_service = GraphRAGService(org_id=self.org_id)
        self.document_processor = DocumentProcessor()
        
        # 使用增强解析器（集成MinerU）
        try:
            self.enhanced_parser = EnhancedDocumentParser(use_mineru=True)
            logger.info("Enhanced document parser (with MinerU) initialized")
        except Exception as e:
            logger.warning(f"Failed to initialize enhanced parser: {e}, using basic parser")
            self.enhanced_parser = None
        
        # 统计信息
        self.stats = {
            "excel_files": 0,
            "pdf_files": 0,
            "docx_files": 0,
            "ppt_files": 0,
            "audio_files": 0,
            "total_documents": 0,
            "total_entities": 0,
            "total_relations": 0,
            "errors": [],
        }
    
    async def ingest_all(self) -> Dict[str, Any]:
        """
        摄入所有数据
        
        Returns:
            摄入统计信息
        """
        logger.info(f"Starting data ingestion from: {DATA_DIR}")
        
        if not DATA_DIR.exists():
            logger.error(f"Data directory not found: {DATA_DIR}")
            return self.stats
        
        # 1. 摄入产品权益（Excel）
        await self._ingest_product_benefits()
        
        # 2. 摄入SOP和话术（PDF、Word、PPT）
        await self._ingest_sop_and_scripts()
        
        # 3. 摄入销售冠军经验（Word）
        await self._ingest_champion_experience()
        
        # 4. 处理销售录音（音频，可选）
        # await self._ingest_sales_recordings()  # 需要STT，暂时跳过
        
        # 5. 重建社区（GraphRAG）
        logger.info("Rebuilding communities...")
        await self.graph_rag_service.rebuild_communities(levels=2)
        
        logger.info(f"Ingestion completed. Stats: {self.stats}")
        return self.stats
    
    async def _ingest_product_benefits(self):
        """摄入产品权益数据（Excel）"""
        product_dir = DATA_DIR / "产品权益"
        if not product_dir.exists():
            logger.warning(f"Product benefits directory not found: {product_dir}")
            return
        
        logger.info("Ingesting product benefits (Excel files)...")
        
        # 处理主目录的Excel文件
        for excel_file in product_dir.glob("*.xlsx"):
            if excel_file.name.startswith("~$"):  # 跳过临时文件
                continue
            
            try:
                await self._ingest_excel_file(excel_file, doc_type="product_benefit")
                self.stats["excel_files"] += 1
            except Exception as e:
                logger.error(f"Failed to ingest {excel_file}: {e}")
                self.stats["errors"].append(f"{excel_file}: {str(e)}")
        
        # 处理竞品分析
        competitor_dir = product_dir / "竞品分析"
        if competitor_dir.exists():
            for excel_file in competitor_dir.glob("*.xlsx"):
                if excel_file.name.startswith("~$"):
                    continue
                
                try:
                    await self._ingest_excel_file(excel_file, doc_type="competitor_analysis")
                    self.stats["excel_files"] += 1
                except Exception as e:
                    logger.error(f"Failed to ingest {excel_file}: {e}")
                    self.stats["errors"].append(f"{excel_file}: {str(e)}")
    
    async def _ingest_sop_and_scripts(self):
        """摄入SOP和话术数据（PDF、Word、PPT）"""
        sop_dir = DATA_DIR / "销售成交营销SOP和话术"
        if not sop_dir.exists():
            logger.warning(f"SOP directory not found: {sop_dir}")
            return
        
        logger.info("Ingesting SOP and scripts (PDF, Word, PPT)...")
        
        # PDF文件
        for pdf_file in sop_dir.glob("*.pdf"):
            try:
                await self._ingest_pdf_file(pdf_file, doc_type="sop")
                self.stats["pdf_files"] += 1
            except Exception as e:
                logger.error(f"Failed to ingest {pdf_file}: {e}")
                self.stats["errors"].append(f"{pdf_file}: {str(e)}")
        
        # Word文件
        for docx_file in sop_dir.glob("*.docx"):
            try:
                await self._ingest_docx_file(docx_file, doc_type="script")
                self.stats["docx_files"] += 1
            except Exception as e:
                logger.error(f"Failed to ingest {docx_file}: {e}")
                self.stats["errors"].append(f"{docx_file}: {str(e)}")
        
        # DOC文件
        for doc_file in sop_dir.glob("*.doc"):
            try:
                await self._ingest_doc_file(doc_file, doc_type="script")
                self.stats["docx_files"] += 1
            except Exception as e:
                logger.error(f"Failed to ingest {doc_file}: {e}")
                self.stats["errors"].append(f"{doc_file}: {str(e)}")
        
        # PPT文件（转换为文本）
        for ppt_file in sop_dir.glob("*.ppt*"):
            try:
                await self._ingest_ppt_file(ppt_file, doc_type="training")
                self.stats["ppt_files"] += 1
            except Exception as e:
                logger.error(f"Failed to ingest {ppt_file}: {e}")
                self.stats["errors"].append(f"{ppt_file}: {str(e)}")
    
    async def _ingest_champion_experience(self):
        """摄入销售冠军经验（Word）"""
        experience_dir = DATA_DIR / "销售冠军成交经验分享"
        if not experience_dir.exists():
            logger.warning(f"Experience directory not found: {experience_dir}")
            return
        
        logger.info("Ingesting champion experience (Word)...")
        
        for docx_file in experience_dir.glob("*.docx"):
            try:
                await self._ingest_docx_file(docx_file, doc_type="case")
                self.stats["docx_files"] += 1
            except Exception as e:
                logger.error(f"Failed to ingest {docx_file}: {e}")
                self.stats["errors"].append(f"{docx_file}: {str(e)}")
    
    async def _ingest_excel_file(self, file_path: Path, doc_type: str = "knowledge"):
        """摄入Excel文件（使用增强解析器）"""
        logger.info(f"Processing Excel: {file_path.name}")
        
        try:
            # 使用增强解析器
            if self.enhanced_parser:
                parsed = self.enhanced_parser.parse_excel(file_path)
                content = parsed["text"]
                metadata = parsed["metadata"]
            else:
                # 降级到基本解析
                if not HAS_EXCEL:
                    logger.warning("openpyxl not installed, skipping Excel files")
                    return
                
                import openpyxl
                workbook = openpyxl.load_workbook(file_path, data_only=True)
                text_parts = []
                
                for sheet_name in workbook.sheetnames:
                    sheet = workbook[sheet_name]
                    text_parts.append(f"【{sheet_name}】\n")
                    
                    for row in sheet.iter_rows(values_only=True):
                        row_text = "\t".join(str(cell) if cell else "" for cell in row)
                        if row_text.strip():
                            text_parts.append(row_text)
                    text_parts.append("\n")
                
                content = "\n".join(text_parts)
                metadata = {
                    "sheet_count": len(workbook.sheetnames),
                    "parser": "openpyxl",
                }
            
            if not content.strip():
                logger.warning(f"Empty Excel file: {file_path.name}")
                return
            
            # 生成文档ID
            doc_id = f"excel_{uuid.uuid4().hex[:8]}"
            
            # 摄入到向量数据库
            await self._ingest_to_vector_db(
                doc_id=doc_id,
                content=content,
                filename=file_path.name,
                doc_type=doc_type,
                metadata={
                    "source": str(file_path),
                    "file_type": "excel",
                    **metadata,
                },
            )
            
            # 摄入到GraphRAG
            await self._ingest_to_graph_rag(
                doc_id=doc_id,
                content=content,
                doc_type=doc_type,
            )
            
            self.stats["total_documents"] += 1
            
        except Exception as e:
            logger.error(f"Failed to ingest Excel {file_path.name}: {e}")
            self.stats["errors"].append(f"{file_path}: {str(e)}")
            raise
    
    async def _ingest_pdf_file(self, file_path: Path, doc_type: str = "knowledge"):
        """摄入PDF文件（使用增强解析器）"""
        logger.info(f"Processing PDF: {file_path.name}")
        
        try:
            # 使用增强解析器（优先MinerU）
            if self.enhanced_parser:
                parsed = self.enhanced_parser.parse_pdf(file_path)
                content = parsed["text"]
                metadata = parsed["metadata"]
                
                # 如果有Markdown，也保存
                if parsed.get("markdown"):
                    metadata["has_markdown"] = True
                
                # 如果有表格，添加到元数据
                if parsed.get("tables"):
                    metadata["table_count"] = len(parsed["tables"])
            else:
                # 降级到基本解析
                if not HAS_PDF:
                    logger.warning("No PDF parser available, skipping")
                    return
                
                pdf_doc = fitz.open(str(file_path))
                text_parts = []
                for page_num in range(len(pdf_doc)):
                    page = pdf_doc[page_num]
                    text = page.get_text()
                    if text.strip():
                        text_parts.append(text)
                pdf_doc.close()
                content = "\n\n".join(text_parts)
                metadata = {
                    "source": str(file_path),
                    "file_type": "pdf",
                    "page_count": len(pdf_doc),
                    "parser": "pymupdf",
                }
            
            if not content.strip():
                logger.warning(f"Empty PDF file: {file_path.name}")
                return
            
            # 生成文档ID
            doc_id = f"pdf_{uuid.uuid4().hex[:8]}"
            
            # 摄入到向量数据库
            await self._ingest_to_vector_db(
                doc_id=doc_id,
                content=content,
                filename=file_path.name,
                doc_type=doc_type,
                metadata={
                    "source": str(file_path),
                    **metadata,
                },
            )
            
            # 摄入到GraphRAG
            await self._ingest_to_graph_rag(
                doc_id=doc_id,
                content=content,
                doc_type=doc_type,
            )
            
            self.stats["total_documents"] += 1
            
        except Exception as e:
            logger.error(f"Failed to ingest PDF {file_path.name}: {e}")
            self.stats["errors"].append(f"{file_path}: {str(e)}")
            raise
    
    async def _ingest_docx_file(self, file_path: Path, doc_type: str = "knowledge"):
        """摄入Word文档（.docx，使用增强解析器）"""
        logger.info(f"Processing DOCX: {file_path.name}")
        
        try:
            # 使用增强解析器
            if self.enhanced_parser:
                parsed = self.enhanced_parser.parse_docx(file_path)
                content = parsed["text"]
                metadata = parsed["metadata"]
            else:
                # 降级到基本解析
                if not HAS_DOCX:
                    logger.warning("python-docx not installed, skipping DOCX files")
                    return
                
                doc = Document(str(file_path))
                text_parts = []
                
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_parts.append(paragraph.text)
                
                for table in doc.tables:
                    for row in table.rows:
                        row_text = "\t".join(cell.text for cell in row.cells)
                        if row_text.strip():
                            text_parts.append(row_text)
                
                content = "\n".join(text_parts)
                metadata = {
                    "parser": "python-docx",
                }
        
            if not content.strip():
                logger.warning(f"Empty DOCX file: {file_path.name}")
                return
            
            # 生成文档ID
            doc_id = f"docx_{uuid.uuid4().hex[:8]}"
            
            # 摄入到向量数据库
            await self._ingest_to_vector_db(
                doc_id=doc_id,
                content=content,
                filename=file_path.name,
                doc_type=doc_type,
                metadata={
                    "source": str(file_path),
                    "file_type": "docx",
                    **metadata,
                },
            )
            
            # 摄入到GraphRAG
            await self._ingest_to_graph_rag(
                doc_id=doc_id,
                content=content,
                doc_type=doc_type,
            )
            
            self.stats["total_documents"] += 1
            
        except Exception as e:
            logger.error(f"Failed to ingest DOCX {file_path.name}: {e}")
            self.stats["errors"].append(f"{file_path}: {str(e)}")
            raise
    
    async def _ingest_doc_file(self, file_path: Path, doc_type: str = "knowledge"):
        """摄入Word文档（.doc，旧格式）"""
        # .doc文件需要特殊处理，这里简化处理
        logger.warning(f".doc files require special handling: {file_path.name}")
        # 可以尝试使用python-docx2txt或其他库
        # 暂时跳过
        pass
    
    async def _ingest_ppt_file(self, file_path: Path, doc_type: str = "knowledge"):
        """摄入PPT文件"""
        logger.warning(f"PPT files require special handling: {file_path.name}")
        # PPT文件需要python-pptx库
        # 暂时跳过
        pass
    
    async def _ingest_to_vector_db(
        self,
        doc_id: str,
        content: str,
        filename: str,
        doc_type: str,
        metadata: Dict[str, Any],
    ):
        """摄入到向量数据库"""
        try:
            # 使用文档处理器处理文档
            result = await self.knowledge_service.add_document_with_processing(
                content=content.encode("utf-8"),
                filename=filename,
                content_type="text/plain",
                meta={
                    **metadata,
                    "doc_id": doc_id,
                    "type": doc_type,
                    "org_id": self.org_id,
                },
                doc_type=doc_type,
            )
            
            logger.info(f"Added to vector DB: {doc_id}, chunks: {result.get('chunks_added', 0)}")
        except Exception as e:
            logger.error(f"Failed to add to vector DB: {e}")
            raise
    
    async def _ingest_to_graph_rag(
        self,
        doc_id: str,
        content: str,
        doc_type: str,
    ):
        """摄入到GraphRAG知识图谱"""
        try:
            result = await self.graph_rag_service.ingest_document(
                doc_id=doc_id,
                text=content,
                metadata={
                    "doc_type": doc_type,
                    "org_id": self.org_id,
                },
            )
            
            self.stats["total_entities"] += result.get("total_entities", 0)
            self.stats["total_relations"] += result.get("total_triples", 0)
            
            logger.info(
                f"Added to GraphRAG: {doc_id}, "
                f"entities: {result.get('total_entities', 0)}, "
                f"relations: {result.get('total_triples', 0)}"
            )
        except Exception as e:
            logger.error(f"Failed to add to GraphRAG: {e}")
            # 不抛出异常，允许继续处理其他文档
    
    def print_statistics(self):
        """打印统计信息"""
        print("\n" + "="*60)
        print("数据摄入统计")
        print("="*60)
        print(f"Excel文件: {self.stats['excel_files']}")
        print(f"PDF文件: {self.stats['pdf_files']}")
        print(f"Word文件: {self.stats['docx_files']}")
        print(f"PPT文件: {self.stats['ppt_files']}")
        print(f"音频文件: {self.stats['audio_files']}")
        print(f"总文档数: {self.stats['total_documents']}")
        print(f"总实体数: {self.stats['total_entities']}")
        print(f"总关系数: {self.stats['total_relations']}")
        
        if self.stats['errors']:
            print(f"\n错误数: {len(self.stats['errors'])}")
            for error in self.stats['errors'][:5]:  # 只显示前5个错误
                print(f"  - {error}")
        
        print("="*60 + "\n")


async def main():
    """主函数"""
    ingester = SalesDataIngester(org_id="public")
    
    try:
        stats = await ingester.ingest_all()
        ingester.print_statistics()
        
        # 打印GraphRAG统计
        graph_stats = ingester.graph_rag_service.get_statistics()
        print("\nGraphRAG统计:")
        print(f"  节点数: {graph_stats.get('graph', {}).get('total_nodes', 0)}")
        print(f"  边数: {graph_stats.get('graph', {}).get('total_edges', 0)}")
        print(f"  社区数: {graph_stats.get('communities', {}).get('total_communities', 0)}")
        
    except Exception as e:
        logger.error(f"Data ingestion failed: {e}", exc_info=True)
        raise


if __name__ == "__main__":
    asyncio.run(main())

